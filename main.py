import json
import os
import re
from io import BytesIO
from pathlib import Path
from typing import Literal

import httpx
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from dotenv import load_dotenv
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.responses import HTMLResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field
from pypdf import PdfReader
from pypdf.errors import PdfReadError

load_dotenv()

APP_VERSION = "2.3.0"
BASE_DIR = Path(__file__).resolve().parent
INDEX_FILE = BASE_DIR / "index.html"
STATIC_DIR = BASE_DIR / "static"
MAX_UPLOAD_BYTES = 12 * 1024 * 1024
MAX_CONTEXT_CHARS = 12_000

DEMO_MODE = os.getenv("DEMO_MODE", "true").lower() in {"1", "true", "yes", "on"}
AI_FALLBACK_TO_DEMO = os.getenv("AI_FALLBACK_TO_DEMO", "true").lower() in {"1", "true", "yes", "on"}
LLM_API_KEY = os.getenv("LLM_API_KEY", "")
LLM_BASE_URL = os.getenv("LLM_BASE_URL", "https://api.openai.com/v1").rstrip("/")
LLM_MODEL = os.getenv("LLM_MODEL", "gpt-4.1-mini")

Language = Literal["en", "ha", "yo", "ig", "sw", "zu"]
SUPPORTED_LANGUAGES = ("en", "ha", "yo", "ig", "sw", "zu")
LANGUAGE_NAMES = {"en": "English", "ha": "Hausa", "yo": "Yorùbá", "ig": "Igbo", "sw": "Kiswahili", "zu": "isiZulu"}

LOCALES: dict[str, dict] = {
    "en": {
        "roles": [("Research supervisor", "Problem & impact"), ("Methodology examiner", "Methodology"), ("Evidence reviewer", "Results & evidence"), ("External examiner", "Limitations"), ("Impact reviewer", "Practical application")],
        "questions": ["What problem does your research on '{topic}' solve, and who benefits most from the result?", "Why did you choose this methodology, and what alternative approach did you consider but reject?", "What is the most important result from your work, and what evidence best supports it?", "What is the biggest limitation of your research, and how should that limitation affect interpretation of the results?", "If you turned this research into a real-world solution, what would you do next and how would you measure success?"],
        "labels": {"clarity": "clarity", "relevance": "relevance to the question", "evidence": "use of evidence", "structure": "answer structure"},
        "strength": "Your strongest area in this response is {label}.", "improvement": "Your biggest opportunity for the next answer is {label}.",
        "tips": {"clarity": "Lead with one sentence that directly answers the question before adding detail.", "relevance": "Echo the key part of the examiner's question and connect every point back to it.", "evidence": "Use one concrete result, number, example, or observation that is actually present in your research context.", "structure": "Use a simple structure: claim → evidence → significance → limitation."},
        "feedback": "Your response was assessed for clarity, relevance, evidence and structure. When research context is provided, the coaching also checks whether your answer connects back to that context. This is a practice estimate, not scientific verification.",
        "framework": "A stronger answer structure you can fill with your real research:\n1. Main claim: [direct answer to the examiner's question].\n2. Evidence: [specific result, number, or observation from your study].\n3. Significance: [why that evidence matters].\n4. Limitation: [what the result cannot prove or where caution is needed].",
        "fallback": "The AI provider was unavailable, so local demo coaching was used to keep the practice session working.",
    },
    "ha": {
        "roles": [("Mai kula da bincike", "Matsala da tasiri"), ("Mai nazarin hanya", "Hanyar bincike"), ("Mai duba hujja", "Sakamako da hujja"), ("Mai jarrabawa na waje", "Iyakoki"), ("Mai duba tasiri", "Amfani a aikace")],
        "questions": ["Wace matsala bincikenka kan '{topic}' yake warwarewa, kuma wa zai fi amfana da sakamakonsa?", "Me ya sa ka zaɓi wannan hanyar bincike, kuma wace hanya ce ka yi la'akari da ita amma ba ka yi amfani da ita ba?", "Mene ne mafi muhimmancin sakamako daga aikinka, kuma wace shaida ce ta fi goyon bayansa?", "Mene ne babban iyakar bincikenka, kuma ta yaya wannan iyakar ke shafar fassarar sakamakon?", "Idan za ka mayar da wannan bincike zuwa mafita ta zahiri, mene ne mataki na gaba kuma yaya za ka auna nasara?"],
        "labels": {"clarity": "bayyananniyar amsa", "relevance": "dacewa da tambaya", "evidence": "amfani da hujja", "structure": "tsarin amsa"},
        "strength": "Mafi ƙarfin ɓangaren amsarka shi ne {label}.", "improvement": "Abin da ya fi dacewa ka inganta a amsa ta gaba shi ne {label}.",
        "tips": {"clarity": "Fara da jimla ɗaya da ke ba da amsa kai tsaye kafin ƙarin bayani.", "relevance": "Danganta kowane batu da ainihin tambayar mai jarrabawa.", "evidence": "Yi amfani da sakamako, adadi ko misali da yake cikin bayanan bincikenka na gaske.", "structure": "Yi amfani da tsari mai sauƙi: batu → hujja → muhimmanci → iyaka."},
        "feedback": "An auna amsarka ta fuskar bayyanawa, dacewa, hujja da tsari. Idan ka bayar da bayanan bincike, horon yana duba ko amsarka ta danganta da su. Wannan kimantawar atisaye ce, ba tabbatar da bincike ba.",
        "framework": "Tsarin da za ka iya amfani da shi:\n1. Babban batu: [amsa kai tsaye ga tambayar].\n2. Hujja: [takamaiman sakamako, adadi ko misali daga bincikenka].\n3. Muhimmanci: [me wannan sakamakon yake nufi].\n4. Iyakar bincike: [abin da sakamakon ba zai iya tabbatarwa ba].",
        "fallback": "Ba a samu AI provider ba, don haka an yi amfani da local demo coaching domin kada atisayen ya tsaya.",
    },
    "yo": {
        "roles": [("Olùtọ́jú ìwádìí", "Ìṣòro àti ipa"), ("Olùdánwò ọ̀nà ìwádìí", "Ọ̀nà ìwádìí"), ("Olùṣàyẹ̀wò ẹ̀rí", "Àbájáde àti ẹ̀rí"), ("Olùdánwò ita", "Ààlà ìwádìí"), ("Olùṣàyẹ̀wò ipa", "Ìlò ní ayé gidi")],
        "questions": ["Ìṣòro wo ni ìwádìí rẹ lórí '{topic}' ń yanjú, ta sì ni yóò jèrè jù lọ nínú àbájáde rẹ?", "Kí ló dé tí o fi yan ọ̀nà ìwádìí yìí, ọ̀nà míì wo ni o sì ronú lé lórí ṣùgbọ́n tí o kò lò?", "Kí ni àbájáde tó ṣe pàtàkì jù lọ nínú iṣẹ́ rẹ, ẹ̀rí wo sì ni ó ṣe atilẹyin rẹ jù lọ?", "Kí ni ààlà pàtàkì jù lọ nínú ìwádìí rẹ, báwo ni ààlà náà ṣe yẹ kí ó ní ipa lórí ìtumọ̀ àbájáde?", "Tí o bá fẹ́ yí ìwádìí yìí padà sí ojútùú gidi, kí ni ìgbésẹ̀ tó kàn, báwo ni o sì ṣe máa wọn àṣeyọrí?"],
        "labels": {"clarity": "kíkedere", "relevance": "ìbáṣepọ̀ pẹ̀lú ìbéèrè", "evidence": "lílò ẹ̀rí", "structure": "ètò ìdáhùn"},
        "strength": "Agbára tó ga jù lọ nínú ìdáhùn rẹ ni {label}.", "improvement": "Ohun tó yẹ kí o dojú kọ jù lọ ní ìdáhùn tó kàn ni {label}.",
        "tips": {"clarity": "Bẹ̀rẹ̀ pẹ̀lú gbólóhùn kan tó dá ìbéèrè lóhùn taara.", "relevance": "So gbogbo kókó rẹ mọ́ ìbéèrè olùdánwò.", "evidence": "Lo àbájáde, nọ́mbà tàbí àkíyèsí tó wà nínú ìwádìí rẹ gan-an.", "structure": "Lo ètò: ìdáhùn → ẹ̀rí → ìtumọ̀ → ààlà."},
        "feedback": "A ṣe àyẹ̀wò ìdáhùn rẹ lórí kíkedere, ìbáṣepọ̀, ẹ̀rí àti ètò. Tí àlàyé ìwádìí bá wà, coaching tún wo bóyá ìdáhùn rẹ bá a mu. Èyí jẹ́ ìṣírò ìdánwò, kì í ṣe ìjẹ́rìí sáyẹ́ǹsì.",
        "framework": "Ètò ìdáhùn tó lágbára:\n1. Kókó pàtàkì: [ìdáhùn taara].\n2. Ẹ̀rí: [àbájáde, nọ́mbà tàbí àkíyèsí gidi].\n3. Ìtumọ̀: [ìdí tí ẹ̀rí náà fi ṣe pàtàkì].\n4. Ààlà: [ohun tí àbájáde kò lè fi hàn].",
        "fallback": "Olùpèsè AI kò sí ní àkókò yìí, nítorí náà a lo demo coaching agbègbè kí ìdánwò rẹ má bà a dá dúró.",
    },
    "ig": {
        "roles": [("Onye nlekọta nyocha", "Nsogbu na mmetụta"), ("Onye nyocha usoro", "Usoro nyocha"), ("Onye nyochaa ihe akaebe", "Nsonaazụ na ihe akaebe"), ("Onye nyocha mpụga", "Oke nyocha"), ("Onye nyochaa mmetụta", "Ojiji n'ezi ndụ")],
        "questions": ["Kedu nsogbu nyocha gị gbasara '{topic}' na-edozi, onye kwa ga-erite uru kachasị na nsonaazụ ya?", "Gịnị mere i ji họrọ usoro nyocha a, kedu ụzọ ọzọ i tụlere ma hapụ?", "Kedu nsonaazụ kacha mkpa n'ọrụ gị, kedu ihe akaebe kacha akwado ya?", "Kedu oke kachasị mkpa nke nyocha gị, olee otú oke ahụ kwesịrị isi metụta nkọwa nsonaazụ?", "Ọ bụrụ na ị gbanwee nyocha a ka ọ bụrụ ngwọta n'ezi ndụ, gịnị bụ nzọụkwụ ọzọ, olee otú ị ga-esi tụọ ihe ịga nke ọma?"],
        "labels": {"clarity": "ido anya", "relevance": "ịza ihe a jụrụ", "evidence": "iji ihe akaebe", "structure": "usoro azịza"},
        "strength": "Akụkụ kacha sie ike n'azịza gị bụ {label}.", "improvement": "Ihe kacha mkpa ị ga-emezi n'azịza ọzọ bụ {label}.",
        "tips": {"clarity": "Malite na otu ahịrịokwu na-aza ajụjụ ahụ ozugbo.", "relevance": "Jikọta isi okwu ọ bụla na ajụjụ onye nyocha.", "evidence": "Jiri nsonaazụ, ọnụọgụ ma ọ bụ nchọpụta dị na nyocha gị n'ezie.", "structure": "Jiri usoro: nkwupụta → ihe akaebe → ihe ọ pụtara → oke."},
        "feedback": "A tụlere azịza gị n'ido anya, ịdị mkpa, ihe akaebe na usoro. Ọ bụrụ na e nyere nkọwa nyocha, coaching na-elekwa ma azịza ahụ dabara na ya. Nke a bụ atụmatụ omume, ọ bụghị nkwenye sayensị.",
        "framework": "Usoro azịza ka mma:\n1. Isi nkwupụta: [azịza ozugbo].\n2. Ihe akaebe: [nsonaazụ, ọnụọgụ ma ọ bụ nchọpụta].\n3. Ihe ọ pụtara: [ihe mere ihe akaebe ji dị mkpa].\n4. Oke: [ihe nsonaazụ ahụ na-enweghị ike igosi].",
        "fallback": "Onye na-enye AI adịghị, ya mere ejiri local demo coaching mee ka mmemme ahụ gaa n'ihu.",
    },
    "sw": {
        "roles": [("Msimamizi wa utafiti", "Tatizo na athari"), ("Mtahini wa mbinu", "Mbinu za utafiti"), ("Mkaguzi wa ushahidi", "Matokeo na ushahidi"), ("Mtahini wa nje", "Mipaka"), ("Mkaguzi wa athari", "Matumizi halisi")],
        "questions": ["Utafiti wako kuhusu '{topic}' unatatua tatizo gani, na nani atanufaika zaidi na matokeo yake?", "Kwa nini ulichagua mbinu hii ya utafiti, na ni njia gani mbadala uliyoifikiria lakini ukaiacha?", "Ni matokeo gani muhimu zaidi katika kazi yako, na ni ushahidi gani unaoyaunga mkono zaidi?", "Ni upungufu gani mkubwa zaidi wa utafiti wako, na unapaswa kuathirije tafsiri ya matokeo?", "Kama ungegeuza utafiti huu kuwa suluhisho la matumizi halisi, hatua inayofuata ingekuwa ipi na ungepimaje mafanikio?"],
        "labels": {"clarity": "uwazi", "relevance": "uhusiano na swali", "evidence": "matumizi ya ushahidi", "structure": "muundo wa jibu"},
        "strength": "Eneo lenye nguvu zaidi katika jibu lako ni {label}.", "improvement": "Eneo muhimu zaidi la kuboresha katika jibu linalofuata ni {label}.",
        "tips": {"clarity": "Anza na sentensi moja inayojibu swali moja kwa moja.", "relevance": "Unganisha kila hoja na swali la mtahini.", "evidence": "Tumia matokeo, namba au uchunguzi uliopo katika utafiti wako halisi.", "structure": "Tumia muundo: dai → ushahidi → umuhimu → kikomo."},
        "feedback": "Jibu lako limepimwa kwa uwazi, uhusiano, ushahidi na muundo. Ikiwa muktadha wa utafiti umetolewa, coaching pia huangalia kama jibu linahusiana nao. Hii ni makadirio ya mazoezi, si uthibitisho wa kisayansi.",
        "framework": "Muundo bora wa jibu:\n1. Dai kuu: [jibu la moja kwa moja].\n2. Ushahidi: [matokeo, namba au uchunguzi maalum].\n3. Umuhimu: [kwa nini ushahidi huo ni muhimu].\n4. Kikomo: [kile ambacho matokeo hayawezi kuthibitisha].",
        "fallback": "Mtoa huduma wa AI hakupatikana, kwa hiyo local demo coaching imetumika ili mazoezi yako yaendelee.",
    },
    "zu": {
        "roles": [("Umqondisi wocwaningo", "Inkinga nomthelela"), ("Umhloli wezindlela", "Indlela yocwaningo"), ("Umhloli wobufakazi", "Imiphumela nobufakazi"), ("Umhloli wangaphandle", "Imikhawulo"), ("Umhloli womthelela", "Ukusetshenziswa empilweni")],
        "questions": ["Ucwaningo lwakho ngo-'{topic}' luxazulula yiphi inkinga, futhi ubani ozohlomula kakhulu emiphumeleni yalo?", "Kungani ukhethe le ndlela yocwaningo, futhi iyiphi enye indlela oyicabangile kodwa wangayisebenzisa?", "Yimuphi umphumela obaluleke kakhulu emsebenzini wakho, futhi yibuphi ubufakazi obuwusekela kakhulu?", "Yimuphi umkhawulo omkhulu wocwaningo lwakho, futhi lowo mkhawulo kufanele uthinte kanjani ukuhunyushwa kwemiphumela?", "Uma ungaguqula lolu cwaningo lube yisixazululo sangempela, yisiphi isinyathelo esilandelayo futhi ungalinganisa kanjani impumelelo?"],
        "labels": {"clarity": "ukucaca", "relevance": "ukuhambisana nombuzo", "evidence": "ukusebenzisa ubufakazi", "structure": "isakhiwo sempendulo"},
        "strength": "Ingxenye enamandla kakhulu empendulweni yakho ngu-{label}.", "improvement": "Into ebaluleke kakhulu ongayithuthukisa empendulweni elandelayo ngu-{label}.",
        "tips": {"clarity": "Qala ngomusho owodwa ophendula umbuzo ngqo.", "relevance": "Xhumanisa wonke amaphuzu akho nombuzo womhloli.", "evidence": "Sebenzisa umphumela, inombolo noma okubonile okukhona ngempela ocwaningweni lwakho.", "structure": "Sebenzisa isakhiwo: iphuzu → ubufakazi → ukubaluleka → umkhawulo."},
        "feedback": "Impendulo ihlolwe ngokucaca, ukuhambisana, ubufakazi nesakhiwo. Uma unikeze umongo wocwaningo, coaching ibheka nokuthi impendulo ixhumene nawo. Lokhu ukuqagela kokuzilolonga, hhayi ukuqinisekiswa kwesayensi.",
        "framework": "Uhlaka lwempendulo oluqinile:\n1. Iphuzu elikhulu: [impendulo eqondile].\n2. Ubufakazi: [umphumela, inombolo noma okubonile].\n3. Ukubaluleka: [kungani ubufakazi bubalulekile].\n4. Umkhawulo: [okungeke kufakazelwe umphumela].",
        "fallback": "Umhlinzeki we-AI ubengatholakali, ngakho kusetshenziswe local demo coaching ukuze ukuzilolonga kuqhubeke.",
    },
}

app = FastAPI(title="Final Defence Coach", version=APP_VERSION)
if STATIC_DIR.exists():
    app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")


class ThesisInput(BaseModel):
    topic: str = Field(min_length=1, max_length=300)
    abstract: str = Field(default="", max_length=MAX_CONTEXT_CHARS)
    language: Language = "en"


class AnswerInput(BaseModel):
    topic: str = Field(min_length=1, max_length=300)
    abstract: str = Field(default="", max_length=MAX_CONTEXT_CHARS)
    question: str = Field(min_length=1, max_length=1200)
    answer: str = Field(min_length=1, max_length=6000)
    language: Language = "en"


def language_name(language: Language) -> str:
    return LANGUAGE_NAMES[language]


def panel_templates(language: Language) -> list[dict[str, str]]:
    return [{"role": role, "category": category} for role, category in LOCALES[language]["roles"]]


def demo_questions(topic: str, language: Language) -> list[dict[str, str]]:
    questions = [template.format(topic=topic) for template in LOCALES[language]["questions"]]
    return [{"id": f"q{i + 1}", **panel, "question": question} for i, (panel, question) in enumerate(zip(panel_templates(language), questions, strict=True))]


def normalize_questions(raw_questions: object, language: Language) -> list[dict[str, str]]:
    if not isinstance(raw_questions, list) or len(raw_questions) != 5:
        raise ValueError("Exactly five questions are required")
    templates = panel_templates(language)
    normalized: list[dict[str, str]] = []
    for i, item in enumerate(raw_questions):
        if isinstance(item, str):
            question, role, category = item.strip(), templates[i]["role"], templates[i]["category"]
        elif isinstance(item, dict):
            question = str(item.get("question") or item.get("text") or "").strip()
            role = str(item.get("role") or templates[i]["role"]).strip()
            category = str(item.get("category") or templates[i]["category"]).strip()
        else:
            raise ValueError("Question format is invalid")
        if not question:
            raise ValueError("Question cannot be empty")
        normalized.append({"id": f"q{i + 1}", "role": role, "category": category, "question": question})
    return normalized


WORD_RE = re.compile(r"[\w'’-]+", re.UNICODE)
STOPWORDS = {"the", "and", "for", "with", "that", "this", "what", "why", "how", "your", "from", "into", "are", "was", "were", "you", "our", "their", "da", "kuma", "wace", "mene", "daga", "don", "kwa", "hii", "hiyo", "nini", "jinsi"}
REASONING_MARKERS = ["because", "therefore", "result", "evidence", "data", "study", "saboda", "sakamako", "shaida", "bincike", "nítorí", "ẹ̀rí", "àbájáde", "ìwádìí", "ihe akaebe", "nsonaazụ", "nnyocha", "kwa sababu", "ushahidi", "matokeo", "utafiti", "ngoba", "ubufakazi", "imiphumela", "ucwaningo"]
STRUCTURE_MARKERS = ["first", "second", "finally", "however", "for example", "na farko", "sannan", "a ƙarshe", "àkọ́kọ́", "ní ìparí", "nke mbụ", "n'ikpeazụ", "kwanza", "hatimaye", "okokuqala", "ekugcineni"]


def meaningful_words(text: str) -> set[str]:
    return {token.lower() for token in WORD_RE.findall(text) if len(token) > 2 and token.lower() not in STOPWORDS}


def numeric_tokens(text: str) -> set[str]:
    return set(re.findall(r"\b\d+(?:\.\d+)?%?\b", text))


def score_dimensions(topic: str, question: str, answer: str, abstract: str = "") -> dict[str, int]:
    words = WORD_RE.findall(answer)
    word_count = len(words)
    sentence_count = max(1, len(re.findall(r"[.!?]+", answer)))
    answer_lower = answer.lower()
    answer_terms = meaningful_words(answer)
    question_overlap = len(answer_terms & meaningful_words(f"{topic} {question}"))
    context_terms = meaningful_words(abstract)
    context_overlap = len(answer_terms & context_terms) if context_terms else 0
    shared_numbers = len(numeric_tokens(answer) & numeric_tokens(abstract)) if abstract else 0
    has_reasoning = any(marker in answer_lower for marker in REASONING_MARKERS)
    has_structure = any(marker in answer_lower for marker in STRUCTURE_MARKERS)

    clarity = 42 + min(34, word_count) + min(10, sentence_count * 2)
    if word_count > 180:
        clarity -= 10
    relevance = 36 + min(28, word_count) + min(30, question_overlap * 6)
    evidence = 36 + min(20, word_count // 2) + (12 if has_reasoning else 0)
    if abstract:
        evidence += min(24, context_overlap * 4) + min(8, shared_numbers * 4)
    else:
        evidence += min(8, word_count // 12)
    structure = 44 + min(20, word_count // 4) + (18 if has_structure else 0) + min(10, sentence_count * 2)

    return {
        "clarity": max(30, min(96, clarity)),
        "relevance": max(30, min(96, relevance)),
        "evidence": max(30, min(96, evidence)),
        "structure": max(30, min(96, structure)),
    }


def demo_evaluation(payload: AnswerInput) -> dict:
    dimensions = score_dimensions(payload.topic, payload.question, payload.answer, payload.abstract)
    locale = LOCALES[payload.language]
    strongest = max(dimensions, key=dimensions.get)
    weakest = min(dimensions, key=dimensions.get)
    return {
        "score": round(sum(dimensions.values()) / len(dimensions)),
        "dimensions": dimensions,
        "strengths": [locale["strength"].format(label=locale["labels"][strongest])],
        "improvements": [locale["improvement"].format(label=locale["labels"][weakest])],
        "feedback": locale["feedback"],
        "improved_answer": locale["framework"],
        "next_tip": locale["tips"][weakest],
        "word_count": len(WORD_RE.findall(payload.answer)),
        "context_used": bool(payload.abstract.strip()),
        "mode": "demo",
        "language": payload.language,
    }


def extract_json(text: str) -> dict:
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\s*```$", "", cleaned)
    return json.loads(cleaned)


async def call_llm(messages: list[dict]) -> dict:
    if not LLM_API_KEY:
        raise HTTPException(status_code=503, detail="LLM_API_KEY is not configured")
    headers = {"Authorization": f"Bearer {LLM_API_KEY}", "Content-Type": "application/json"}
    base_body = {"model": LLM_MODEL, "messages": messages, "temperature": 0.25}
    try:
        async with httpx.AsyncClient(timeout=httpx.Timeout(45.0, connect=10.0)) as client:
            for json_mode in (True, False):
                body = dict(base_body)
                if json_mode:
                    body["response_format"] = {"type": "json_object"}
                response = await client.post(f"{LLM_BASE_URL}/chat/completions", headers=headers, json=body)
                if json_mode and response.status_code in {400, 404, 415, 422}:
                    continue
                response.raise_for_status()
                data = response.json()
                return extract_json(data["choices"][0]["message"]["content"])
    except (httpx.HTTPError, KeyError, IndexError, json.JSONDecodeError) as exc:
        raise HTTPException(status_code=502, detail="The AI provider returned an invalid or unavailable response") from exc
    raise HTTPException(status_code=502, detail="The AI provider did not accept the request format")


def fallback_notice(language: Language) -> str:
    return LOCALES[language]["fallback"]


def suggested_topic_from_text(text: str) -> str:
    for line in text.splitlines():
        candidate = line.strip().lstrip("#").strip()
        if 4 <= len(candidate) <= 180:
            return candidate
    return ""


def select_context(text: str) -> tuple[str, bool]:
    text = text.strip()
    if len(text) <= MAX_CONTEXT_CHARS:
        return text, False
    chunk = MAX_CONTEXT_CHARS // 3
    middle = len(text) // 2
    parts = [text[:chunk], text[max(0, middle - chunk // 2): middle + chunk // 2], text[-chunk:]]
    sampled = "\n\n…\n\n".join(part.strip() for part in parts)
    return sampled[:MAX_CONTEXT_CHARS], True


def extract_docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    return "\n".join(blocks)


@app.middleware("http")
async def add_security_headers(request, call_next):
    response = await call_next(request)
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["Referrer-Policy"] = "no-referrer"
    response.headers["Permissions-Policy"] = "microphone=(self)"
    response.headers["Content-Security-Policy"] = "default-src 'self'; script-src 'self'; style-src 'self'; img-src 'self' data:; connect-src 'self'; font-src 'self'; object-src 'none'; base-uri 'none'; frame-ancestors 'none'; form-action 'self'"
    if request.url.path.startswith("/api/") or request.url.path == "/health":
        response.headers["Cache-Control"] = "no-store"
    return response


@app.get("/", response_class=HTMLResponse)
def home() -> HTMLResponse:
    if not INDEX_FILE.exists():
        raise HTTPException(status_code=500, detail="index.html is missing")
    return HTMLResponse(INDEX_FILE.read_text(encoding="utf-8"))


@app.get("/health")
def health() -> dict:
    return {"status": "ok", "version": APP_VERSION, "mode": "demo" if DEMO_MODE else "llm", "fallback_to_demo": AI_FALLBACK_TO_DEMO, "languages": list(SUPPORTED_LANGUAGES), "upload_types": ["pdf", "docx", "txt", "md"], "research_aware_evaluation": True, "representative_context_sampling": True}


@app.post("/api/extract")
async def extract_research(file: UploadFile = File(...)) -> dict:
    filename = (file.filename or "research").strip()
    suffix = Path(filename).suffix.lower()
    if suffix not in {".pdf", ".docx", ".txt", ".md"}:
        raise HTTPException(status_code=415, detail="Use a PDF, DOCX, TXT, or Markdown file")
    content = await file.read(MAX_UPLOAD_BYTES + 1)
    if len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="File is too large; maximum size is 12 MB")
    try:
        if suffix == ".pdf":
            reader = PdfReader(BytesIO(content))
            text = "\n\n".join((page.extract_text() or "").strip() for page in reader.pages)
        elif suffix == ".docx":
            text = extract_docx_text(content)
        else:
            text = content.decode("utf-8-sig")
    except (PdfReadError, PackageNotFoundError, UnicodeDecodeError, ValueError, KeyError) as exc:
        raise HTTPException(status_code=422, detail="The file could not be read as text") from exc
    text = text.strip()
    if not text:
        raise HTTPException(status_code=422, detail="No extractable text was found in this file")
    context, truncated = select_context(text)
    return {"filename": filename, "text": context, "suggested_topic": suggested_topic_from_text(text), "truncated": truncated, "characters": len(context)}


@app.post("/api/questions")
async def generate_questions(payload: ThesisInput) -> dict:
    topic, abstract = payload.topic.strip(), payload.abstract.strip()
    if not topic:
        raise HTTPException(status_code=422, detail="Topic cannot be empty")
    if DEMO_MODE:
        return {"questions": demo_questions(topic, payload.language), "mode": "demo", "language": payload.language}
    target_language = language_name(payload.language)
    context = abstract or "No thesis context was provided. Base the panel on the topic and ask broadly useful defence questions."
    try:
        result = await call_llm([
            {"role": "system", "content": "You are a realistic but supportive university thesis defence panel. Return JSON only with a 'questions' array of exactly five objects. Each object must contain 'role', 'category', and 'question'. Cover problem/impact, methodology, evidence/results, limitations, and practical application. " + f"Write every visible field in {target_language}. Keep questions concise and grounded in the supplied research context; do not invent findings."},
            {"role": "user", "content": f"Thesis topic: {topic}\n\nResearch context: {context}"},
        ])
        questions = normalize_questions(result.get("questions"), payload.language)
        return {"questions": questions, "mode": "llm", "language": payload.language}
    except (HTTPException, ValueError) as exc:
        if not AI_FALLBACK_TO_DEMO:
            if isinstance(exc, HTTPException):
                raise
            raise HTTPException(status_code=502, detail="The AI panel returned an invalid question format") from exc
        return {"questions": demo_questions(topic, payload.language), "mode": "fallback", "language": payload.language, "notice": fallback_notice(payload.language)}


@app.post("/api/evaluate")
async def evaluate_answer(payload: AnswerInput) -> dict:
    if DEMO_MODE:
        return demo_evaluation(payload)
    target_language = language_name(payload.language)
    research_context = payload.abstract.strip() or "No additional research context was supplied. Do not invent missing evidence."
    try:
        result = await call_llm([
            {"role": "system", "content": "Act as a thesis defence coach. Return JSON only with: object 'dimensions' containing integer 0-100 scores for clarity, relevance, evidence, and structure; array 'strengths' with 1-3 concise items; array 'improvements' with 1-3 concise items; short 'feedback'; 'improved_answer'; and one-sentence 'next_tip'. Compare the student's answer only with the supplied thesis topic and research context. Never invent findings, numbers, citations or facts. If evidence is missing, say so or use placeholders. " + f"Write all coaching text in {target_language}."},
            {"role": "user", "content": f"Thesis topic: {payload.topic}\nResearch context: {research_context}\nExaminer question: {payload.question}\nStudent answer: {payload.answer}"},
        ])
        raw = result["dimensions"]
        dimensions = {key: max(0, min(100, int(raw[key]))) for key in ("clarity", "relevance", "evidence", "structure")}
        strengths = [str(x) for x in result["strengths"]][:3]
        improvements = [str(x) for x in result["improvements"]][:3]
        if not strengths or not improvements:
            raise ValueError("Coaching lists cannot be empty")
        return {"score": round(sum(dimensions.values()) / 4), "dimensions": dimensions, "strengths": strengths, "improvements": improvements, "feedback": str(result["feedback"]), "improved_answer": str(result["improved_answer"]), "next_tip": str(result["next_tip"]), "word_count": len(WORD_RE.findall(payload.answer)), "context_used": bool(payload.abstract.strip()), "mode": "llm", "language": payload.language}
    except (HTTPException, KeyError, TypeError, ValueError) as exc:
        if not AI_FALLBACK_TO_DEMO:
            if isinstance(exc, HTTPException):
                raise
            raise HTTPException(status_code=502, detail="The AI evaluation format was invalid") from exc
        response = demo_evaluation(payload)
        response["mode"] = "fallback"
        response["notice"] = fallback_notice(payload.language)
        return response
