import os
import sys
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

os.environ["DEMO_MODE"] = "true"
os.environ["AI_FALLBACK_TO_DEMO"] = "true"
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from fastapi.testclient import TestClient

from main import app, normalize_questions, score_dimensions

client = TestClient(app)
SUPPORTED = ["en", "ha", "yo", "ig", "sw", "zu"]
FIRST_ROLES = {
    "en": "Research supervisor",
    "ha": "Mai kula da bincike",
    "yo": "Olùtọ́jú ìwádìí",
    "ig": "Onye nlekọta nyocha",
    "sw": "Msimamizi wa utafiti",
    "zu": "Umqondisi wocwaningo",
}


def test_home_page_has_clear_first_use_flow_and_external_assets():
    response = client.get("/")
    assert response.status_code == 200
    assert "Final Defence Coach" in response.text
    assert "Multilingual for Africa" in response.text
    assert "English + Hausa" not in response.text
    assert 'id="guidedDemoBtn"' in response.text
    assert 'id="dropzone"' in response.text
    assert 'id="sessionComplete"' in response.text
    assert 'id="scoreExplanation"' in response.text
    assert 'href="/static/styles.css"' in response.text
    assert 'src="/static/i18n.js"' in response.text
    assert 'src="/static/app.js"' in response.text
    for language, label in {"en": "English", "ha": "Hausa", "yo": "Yorùbá", "ig": "Igbo", "sw": "Kiswahili", "zu": "isiZulu"}.items():
        assert f'<option value="{language}">{label}</option>' in response.text


def test_static_frontend_assets_are_served_and_contain_resilience_logic():
    app_js = client.get("/static/app.js")
    css = client.get("/static/styles.css")
    i18n = client.get("/static/i18n.js")
    assert app_js.status_code == 200
    assert "saveCurrentAnswer" in app_js.text
    assert "previous_score" in app_js.text
    assert "sessionStorage" in app_js.text
    assert ".docx" in app_js.text
    assert css.status_code == 200 and ".dropzone" in css.text
    assert i18n.status_code == 200 and "window.FDC_I18N" in i18n.text


def test_health_reports_research_aware_runtime_capabilities():
    response = client.get("/health")
    assert response.status_code == 200
    assert response.json() == {
        "status": "ok",
        "version": "2.3.0",
        "mode": "demo",
        "fallback_to_demo": True,
        "languages": SUPPORTED,
        "upload_types": ["pdf", "docx", "txt", "md"],
        "research_aware_evaluation": True,
        "representative_context_sampling": True,
    }


def test_security_and_privacy_headers_are_present():
    response = client.post("/api/questions", json={"topic": "Test topic", "abstract": "", "language": "en"})
    assert response.status_code == 200
    assert response.headers["x-content-type-options"] == "nosniff"
    assert response.headers["referrer-policy"] == "no-referrer"
    assert response.headers["cache-control"] == "no-store"
    assert "default-src 'self'" in response.headers["content-security-policy"]


@pytest.mark.parametrize("language", SUPPORTED)
def test_generate_five_role_panel_in_every_supported_language(language):
    response = client.post("/api/questions", json={"topic": "AI crop support", "abstract": "", "language": language})
    assert response.status_code == 200
    data = response.json()
    assert data["mode"] == "demo"
    assert data["language"] == language
    assert len(data["questions"]) == 5
    assert all(set(item) == {"id", "role", "category", "question"} for item in data["questions"])
    assert len({item["role"] for item in data["questions"]}) == 5
    assert data["questions"][0]["role"] == FIRST_ROLES[language]


def test_question_normalizer_accepts_plain_strings_for_provider_compatibility():
    questions = normalize_questions([f"Question {number}" for number in range(1, 6)], "en")
    assert len(questions) == 5
    assert questions[0]["role"] == "Research supervisor"
    assert questions[4]["category"] == "Practical application"


@pytest.mark.parametrize("language", SUPPORTED)
def test_demo_coaching_returns_structure_for_each_language(language):
    response = client.post(
        "/api/evaluate",
        json={
            "topic": "AI crop support",
            "abstract": "The project helps smallholder farmers get an earlier crop disease signal before seeking expert support.",
            "question": "What problem does the project solve?",
            "answer": "The project helps smallholder farmers get an earlier crop disease signal because expert support may not be immediately available.",
            "language": language,
        },
    )
    assert response.status_code == 200
    data = response.json()
    assert data["mode"] == "demo"
    assert data["language"] == language
    assert data["context_used"] is True
    assert 0 <= data["score"] <= 100
    assert set(data["dimensions"]) == {"clarity", "relevance", "evidence", "structure"}
    assert data["strengths"] and data["improvements"] and data["feedback"] and data["improved_answer"] and data["next_tip"]


def test_research_context_can_improve_evidence_score_when_answer_matches_it():
    topic = "AI crop disease detection"
    question = "What evidence supports the approach?"
    answer = "The study observed an earlier crop disease signal for smallholder farmers because the workflow uses mobile images."
    unrelated_context = "This document discusses university library opening hours and student parking."
    related_context = "The study uses mobile images to give smallholder farmers an earlier crop disease signal."
    unrelated = score_dimensions(topic, question, answer, unrelated_context)
    related = score_dimensions(topic, question, answer, related_context)
    assert related["evidence"] > unrelated["evidence"]


def test_cross_language_style_answer_is_not_crushed_by_zero_word_overlap():
    dimensions = score_dimensions(
        "AI crop support",
        "What problem does the project solve?",
        "Mradi huu unasaidia wakulima wadogo kupata ishara ya mapema kabla ya kutafuta msaada wa mtaalamu.",
        "The project gives smallholder farmers an earlier signal before they seek expert support.",
    )
    assert dimensions["relevance"] >= 50


def test_unverified_number_is_not_rewarded_as_research_evidence():
    topic = "AI crop support"
    question = "What evidence supports it?"
    context = "The study describes interviews with farmers but gives no accuracy percentage."
    without_number = score_dimensions(topic, question, "The study describes interviews with farmers because their experience informs the design.", context)
    invented_number = score_dimensions(topic, question, "The study is 99 percent accurate because farmers were interviewed.", context)
    assert invented_number["evidence"] <= without_number["evidence"] + 4


def test_english_framework_does_not_invent_research_facts():
    response = client.post(
        "/api/evaluate",
        json={"topic": "AI crop support", "abstract": "", "question": "What evidence supports your result?", "answer": "The project is intended to give an earlier signal.", "language": "en"},
    )
    assert response.status_code == 200
    data = response.json()
    assert "[specific result" in data["improved_answer"]
    assert data["context_used"] is False


def test_extract_plain_text_research_file():
    response = client.post(
        "/api/extract",
        files={"file": ("thesis.txt", b"AI-assisted crop disease detection\nThis study explores an accessible workflow for smallholder farmers.", "text/plain")},
    )
    assert response.status_code == 200
    data = response.json()
    assert data["filename"] == "thesis.txt"
    assert data["suggested_topic"] == "AI-assisted crop disease detection"
    assert "smallholder farmers" in data["text"]
    assert data["truncated"] is False


def test_extract_docx_research_file():
    document = Document()
    document.add_heading("Community health referral assistant", level=1)
    document.add_paragraph("This thesis explores a low-resource referral workflow for rural clinics.")
    buffer = BytesIO()
    document.save(buffer)
    response = client.post(
        "/api/extract",
        files={"file": ("thesis.docx", buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert response.status_code == 200
    data = response.json()
    assert data["filename"] == "thesis.docx"
    assert data["suggested_topic"] == "Community health referral assistant"
    assert "rural clinics" in data["text"]


def test_long_research_samples_beginning_middle_and_end_instead_of_only_start():
    body = "BEGIN_MARKER\n" + ("A" * 6500) + "\nMIDDLE_MARKER\n" + ("B" * 6500) + "\nEND_MARKER"
    response = client.post("/api/extract", files={"file": ("long.txt", body.encode(), "text/plain")})
    assert response.status_code == 200
    data = response.json()
    assert data["truncated"] is True
    assert data["characters"] <= 12000
    assert "BEGIN_MARKER" in data["text"]
    assert "MIDDLE_MARKER" in data["text"]
    assert "END_MARKER" in data["text"]


def test_extract_rejects_unsupported_file_type():
    response = client.post("/api/extract", files={"file": ("notes.xlsx", b"not supported", "application/octet-stream")})
    assert response.status_code == 415


def test_blank_topic_and_unknown_language_are_rejected():
    blank = client.post("/api/questions", json={"topic": " ", "abstract": "Anything", "language": "en"})
    unknown = client.post("/api/questions", json={"topic": "Test topic", "abstract": "", "language": "xx"})
    assert blank.status_code == 422
    assert unknown.status_code == 422
